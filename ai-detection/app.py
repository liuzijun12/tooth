import os
import torch
import torchvision.transforms as transforms
from flask import Flask, request, jsonify, send_from_directory, session, redirect, url_for, Response
from flask_cors import CORS
from functools import wraps
import logging
import webbrowser
import threading
import time
import requests
from werkzeug.utils import secure_filename
from ultralytics import YOLO
import json
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import threading
import sys

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局配置
MODEL_PATH = "best.pt"  # 训练好的最佳模型
IMAGE_SIZE = 320               # 训练时的图像尺寸
DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# 全局模型变量
model = None

# 大语言模型配置
LLM_CONFIG = {
    "model_type": "local",  # local 或 api
    "model_name": "deepseek-r1:1.5b",  # 默认模型名称
    "api_url": "",  # API URL，如果使用API
    "api_key": "",  # API密钥，如果使用API
}


def input_with_timeout(prompt, timeout=10):
    """
    支持超时的输入函数，超时后返回 None
    """
    user_input = {}

    def get_input():
        try:
            user_input['data'] = input(prompt)
        except Exception:
            user_input['data'] = None

    input_thread = threading.Thread(target=get_input)
    input_thread.daemon = True
    input_thread.start()
    input_thread.join(timeout)

    return user_input.get('data')





def setup_llm_config():
    """设置大语言模型配置"""
    print("\n" + "="*50)
    print("欢迎使用口腔疾病预测中心")
    print("="*50)

    print("\n请选择AI问诊使用的大语言模型类型:")
    print("1. 本地模型 (Ollama)")
    print("2. 网络API (如OpenAI, Azure等)")

    # choice = ""
    # while choice not in ["1", "2"]:
    #     choice = input("\n请输入选项 (1/2): ").strip()

    try:
        choice = input_with_timeout("\n请输入选项 (1/2): ", timeout=10)
        if not choice:
            print("10秒内未输入，默认选择本地模型 (2)")
            choice = "2"
    except OSError:
        print("检测到后台启动（nohup环境），默认选择本地模型 (12")
        choice = "2"

    choice = choice.strip()

    if choice == "1":
        LLM_CONFIG["model_type"] = "local"

        # 获取可用的本地模型列表
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            if response.status_code == 200:
                models = response.json().get("models", [])
                if models:
                    print("\n可用的本地模型:")
                    for i, model in enumerate(models, 1):
                        print(f"{i}. {model.get('name')}")

                    model_choice = ""
                    while not model_choice.isdigit() or int(model_choice) < 1 or int(model_choice) > len(models):
                        model_choice = input(f"\n请选择模型 (1-{len(models)}): ").strip()

                    selected_model = models[int(model_choice)-1].get('name')
                    LLM_CONFIG["model_name"] = selected_model
                    print(f"\n已选择模型: {selected_model}")
                else:
                    print("\n未找到可用的本地模型，将使用默认模型: deepseek-r1:1.5b")
                    print("如需安装模型，请运行: ollama pull deepseek-r1:1.5b")
            else:
                print("\nOllama服务未响应，请确保Ollama正在运行")
                print("将使用默认模型: deepseek-r1:1.5b")
        except Exception as e:
            print(f"\n获取本地模型列表失败: {str(e)}")
            print("将使用默认模型: deepseek-r1:1.5b")
            print("请确保Ollama服务正在运行，或者运行: ollama pull deepseek-r1:1.5b")

    else:  # choice == "2"
        LLM_CONFIG["model_type"] = "api"

        # api_url = input("\n请输入API URL (例如: https://api.openai.com/v1/chat/completions): ").strip()
        api_url = "https://api.deepseek.com"
        # https: // api.deepseek.com / v1
        # api_key = input("请输入API密钥: ").strip()
        api_key = "sk-aa19071022fe47f1aaef8a0215749ddc"
        # model_name = input("请输入模型名称 (例如: gpt-3.5-turbo): ").strip()
        model_name = "deepseek-chat"

        LLM_CONFIG["api_url"] = api_url
        LLM_CONFIG["api_key"] = api_key

        if model_name:
            LLM_CONFIG["model_name"] = model_name

    print("\n配置完成！正在启动服务器...\n")
    return LLM_CONFIG

# 加载模型
def load_model(model_path):
    """加载YOLO模型"""
    try:
        print(f"Loading model from {model_path}")
        model = YOLO(model_path)
        print("Model loaded successfully")
        return model
    except Exception as e:
        print(f"Error loading model: {str(e)}")
        return None

# 预处理函数
def preprocess_image(image):
    transform = transforms.Compose([
        transforms.Resize((IMAGE_SIZE, IMAGE_SIZE)),
        transforms.ToTensor(),
        transforms.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])
    ])
    return transform(image).unsqueeze(0).to(DEVICE)

app = Flask(__name__, static_folder='ai_detection')
CORS(app, supports_credentials=True)  # 支持跨域请求
app.secret_key = 'your-secret-key'  # 更换为随机的密钥
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# 确保上传文件夹存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 加载模型
model = load_model(MODEL_PATH)
if model is None:
    logger.error("Failed to load model. Predictions will not work!")

# 添加 session 配置
app.config.update(
    SESSION_COOKIE_SECURE=False,  # 开发环境设为 False
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax'
)

# 路由：服务静态文件
@app.route('/')
def index():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory(app.static_folder, path)

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            return jsonify({"error": "Please login first"}), 401
        return f(*args, **kwargs)
    return decorated_function

# 预测接口
@app.route("/predict", methods=["POST"])
def predict():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    if model is None:
        return jsonify({'error': 'Model not loaded'}), 500
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    try:
        # 保存上传的文件
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # 进行预测
        results = model(filepath)
        
        # 获取检测结果
        detections = results[0].boxes
        num_detections = len(detections)
        
        # 初始化疾病检测结果
        diseases = {
            'calculus': 0,
            'caries': 0,
            'gingivitis': 0,
            'hypodontia': 0,
            'tooth_discolation': 0,
            'ulcer': 0
        }
        
        # 统计每种疾病的检测结果
        if num_detections > 0:
            for box in detections:
                cls = int(box.cls[0])  # 获取类别索引
                conf = float(box.conf[0])  # 获取置信度
                class_name = results[0].names[cls]  # 获取类别名称
                if class_name in diseases:
                    diseases[class_name] = max(diseases[class_name], conf)  # 保存最高置信度
        
        # 构建响应数据
        response_data = {
            'diseases': diseases,
            'num_detections': num_detections
        }
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Error during prediction: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        # 清理上传的文件
        if os.path.exists(filepath):
            os.remove(filepath)

# 登录接口
@app.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    username = data.get("username")
    password = data.get("password")

    if username == "admin" and password == "123456":
        session['logged_in'] = True
        session['username'] = username
        return jsonify({
            "success": True,
            "message": "登录成功",
            "username": username
        })
    else:
        return jsonify({
            "success": False,
            "message": "用户名或密码错误"
        }), 401

@app.route("/logout", methods=["POST"])
def logout():
    session.clear()
    return jsonify({"success": True, "message": "登出成功"})

@app.route("/chat_ollama", methods=["POST"])
def chat_ollama():
    try:
        data = request.get_json()
        prompt = data.get("prompt")
        stream = data.get("stream", False)
        model_name = data.get("model", LLM_CONFIG["model_name"])

        if not prompt:
            return jsonify({"error": "请输入您的问题"}), 400

        logger.info(f"收到聊天请求: prompt={prompt}, model={model_name}")

        # 构建上下文
        context = f"""你是一位专业的牙科医生，请直接回答问题，不要有思考过程，不要说"您好，我是AI牙科助手"，不要重复用户的问题，直接给出专业的回答。
        
        用户问题：{prompt}"""

        try:
            # 根据配置选择不同的处理方式
            if LLM_CONFIG["model_type"] == "local":
                # 首先测试Ollama服务是否可用
                try:
                    health_check = requests.get("http://localhost:11434/api/tags", timeout=5)
                    if health_check.status_code != 200:
                        logger.error("Ollama服务未运行或无响应")
                        return jsonify({
                            "error": "AI服务未启动",
                            "message": "请确保已安装Ollama并运行服务。\n1. 安装Ollama\n2. 运行: ollama pull deepseek-r1:1.5b\n3. 启动Ollama服务"
                        }), 503
                except requests.exceptions.ConnectionError:
                    return jsonify({
                        "error": "无法连接到Ollama服务",
                        "message": "请确保Ollama服务正在运行，默认端口为11434"
                    }), 503

                # 调用Ollama API
                logger.info(f"正在调用Ollama API，使用模型: {model_name}")
                
                if stream:
                    def generate():
                        response = requests.post(
                            "http://localhost:11434/api/generate",
                            json={
                                "model": model_name,
                                "prompt": context,
                                "stream": True,
                                "temperature": 0.7,
                                "top_p": 0.9
                            },
                            stream=True
                        )
                        
                        for line in response.iter_lines():
                            if line:
                                try:
                                    data = json.loads(line)
                                    if "response" in data:
                                        yield f"data: {json.dumps({'response': data['response']})}\n\n"
                                except json.JSONDecodeError:
                                    continue
                    
                    return Response(generate(), mimetype='text/event-stream')
                else:
                    response = requests.post(
                        "http://localhost:11434/api/generate",
                        json={
                            "model": model_name,
                            "prompt": context,
                            "stream": False,
                            "temperature": 0.7,
                            "top_p": 0.9
                        }
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        return jsonify({"response": result.get("response", "抱歉，我现在无法回答您的问题。")})
                    else:
                        logger.error(f"Ollama API错误响应: {response.text}")
                        return jsonify({
                            "error": "AI服务响应错误",
                            "message": "请检查Ollama服务状态和模型是否正确加载"
                        }), 500

            else:  # LLM_CONFIG["model_type"] == "api"
                if not LLM_CONFIG["api_url"] or not LLM_CONFIG["api_key"]:
                    return jsonify({
                        "error": "API配置不完整",
                        "message": "请提供有效的API URL和API Key"
                    }), 400

                # 调用外部API
                logger.info(f"正在调用外部API: {LLM_CONFIG['api_url']}")
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {LLM_CONFIG['api_key']}"
                }

                payload = {
                    "model": LLM_CONFIG["model_name"],
                    "messages": [
                        {"role": "system", "content": "你是一位专业的牙科医生..."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.7,
                    "max_tokens": 1000
                }

                response = requests.post(
                    f"{LLM_CONFIG['api_url'].rstrip('/')}/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=30
                )

                if response.status_code == 200:
                    result = response.json()
                    ai_response = result["choices"][0]["message"]["content"]
                    return jsonify({"response": ai_response})
                else:
                    logger.error(f"API错误响应: {response.text}")
                    return jsonify({
                        "error": "API服务错误",
                        "message": "请检查API配置是否正确"
                    }), 500

        except requests.exceptions.Timeout:
            logger.error("请求超时")
            return jsonify({
                "error": "请求超时",
                "message": "AI服务响应时间过长，请稍后重试"
            }), 504
        except requests.exceptions.RequestException as e:
            logger.error(f"请求错误: {str(e)}")
            return jsonify({
                "error": "请求错误",
                "message": str(e)
            }), 500

    except Exception as e:
        logger.error(f"处理聊天请求时出错: {str(e)}")
        return jsonify({
            "error": "服务器错误",
            "message": "处理请求时发生错误，请稍后重试"
        }), 500

@app.route("/export_word", methods=["POST"])
def export_word():
    try:
        data = request.get_json()
        detection_results = data.get("detection_results")
        
        if not detection_results:
            return jsonify({"error": "No detection results provided"}), 400
            
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('口腔健康检测报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加时间
        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f'检测时间：{current_time}')
        
        # 添加检测结果
        doc.add_heading('检测结果', level=1)
        doc.add_paragraph(f'检测到的问题数量：{detection_results["num_detections"]}')
        
        # 疾病名称的中文映射
        disease_names = {
            'calculus': '牙结石',
            'caries': '龋齿',
            'gingivitis': '牙龈炎',
            'hypodontia': '牙齿缺失',
            'tooth_discolation': '牙齿变色',
            'ulcer': '溃疡'
        }
        
        # 添加各疾病的检测结果
        doc.add_heading('详细分析', level=2)
        for disease, probability in detection_results['diseases'].items():
            if probability > 0:  # 只显示检测到的问题
                p = doc.add_paragraph()
                p.add_run(f'{disease_names[disease]}: ').bold = True
                p.add_run(f'置信度 {probability:.1%}')
        
        # 添加建议
        doc.add_heading('建议', level=1)
        has_issues = any(prob > 0.5 for prob in detection_results['diseases'].values())
        
        if has_issues:
            suggestions = [
                "请及时就医进行专业检查",
                "保持良好的口腔卫生习惯",
                "定期进行口腔检查",
                "避免过度食用甜食和碳酸饮料"
            ]
        else:
            suggestions = [
                "继续保持良好的口腔卫生习惯",
                "建议每半年进行一次口腔检查",
                "使用合适的牙刷和牙膏",
                "保持健康的饮食习惯"
            ]
            
        for suggestion in suggestions:
            doc.add_paragraph(suggestion, style='List Bullet')
        
        # 保存文档
        reports_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'reports')
        os.makedirs(reports_dir, exist_ok=True)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        report_filename = f'dental_report_{timestamp}.docx'
        report_path = os.path.join(reports_dir, report_filename)
        doc.save(report_path)
        
        # 返回报告URL
        report_url = url_for('download_report', filename=report_filename)
        return jsonify({"report_url": report_url})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Not found"}), 404

@app.errorhandler(500)
def server_error(error):
    return jsonify({"error": "Internal server error"}), 500

# 添加自动打开浏览器的函数
def open_browser():
    time.sleep(1)  # 等待服务器启动
    webbrowser.open('http://127.0.0.1:5000')

if __name__ == "__main__":
    # 设置大语言模型配置
    setup_llm_config()

    # 加载模型
    model = load_model(MODEL_PATH)
    if model is None:
        print("Warning: Running without model. Predictions will not work!")
        logger.error("Failed to load model. Predictions will not work!")

    # 在新线程中打开浏览器
    threading.Thread(target=open_browser).start()

    logger.info("Server started")
    # 启动服务器
    app.run(host="0.0.0.0", port=5000, debug=True, use_reloader=False) 